"""Build the Weather Cup 2026 group-stage report as a DOCX artifact."""

from __future__ import annotations

import argparse
import json
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
DATA_JS = ROOT / "website" / "mvp" / "data.js"

INK = RGBColor(23, 32, 38)
BLUE = RGBColor(32, 74, 120)
MUTED = RGBColor(97, 112, 122)
SURFACE = RGBColor(238, 243, 245)
LINE = RGBColor(216, 225, 229)
TEAL = RGBColor(15, 139, 141)
CORAL = RGBColor(217, 93, 57)


def _load_payload(path: Path) -> dict[str, Any]:
    raw = path.read_text(encoding="utf-8").strip()
    prefix = "window.WM_MVP_DATA = "
    if not raw.startswith(prefix):
        raise ValueError(f"Unsupported data.js format: {path}")
    raw = raw[len(prefix):]
    if raw.endswith(";"):
        raw = raw[:-1]
    return json.loads(raw)


def _set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def _set_cell_border(cell, color: str = "D8E1E5") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right"):
        elem = tc_borders.find(qn(f"w:{edge}"))
        if elem is None:
            elem = OxmlElement(f"w:{edge}")
            tc_borders.append(elem)
        elem.set(qn("w:val"), "single")
        elem.set(qn("w:sz"), "8")
        elem.set(qn("w:color"), color)


def _style_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    title = doc.styles["Title"]
    title.font.name = "Calibri"
    title.font.size = Pt(24)
    title.font.color.rgb = BLUE

    for style_name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, RGBColor(31, 77, 120), 8, 4),
    ):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def _add_run(paragraph, text: str, *, size: int | None = None, color: RGBColor | None = None, bold: bool = False) -> None:
    run = paragraph.add_run(text)
    run.font.name = "Calibri"
    run.bold = bold
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color


def _add_cover(doc: Document, report: dict[str, Any], metadata: dict[str, Any]) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(72)
    p.paragraph_format.space_after = Pt(18)
    _add_run(p, "Weather Cup 2026 Report", size=12, color=TEAL, bold=True)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(4)
    _add_run(title, "Gruppenphase: Der erste Weather-Cup-Report", size=28, color=BLUE, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(18)
    _add_run(subtitle, "Wissenschaftlich belastbare Zwischenbilanz nach 72 Gruppenspielen", size=14, color=MUTED)

    meta = doc.add_table(rows=2, cols=3)
    meta.autofit = False
    widths = [Inches(2.15), Inches(2.15), Inches(2.2)]
    for col, width in zip(meta.columns, widths):
        col.width = width
    labels = ["Datenstand", "Turnier-Scope", "Exportquelle"]
    values = [
        metadata.get("exported_at", "–"),
        report.get("scope_label_de", "–"),
        metadata.get("source", "–"),
    ]
    for idx, label in enumerate(labels):
        cell = meta.cell(0, idx)
        cell.text = label
        _set_cell_shading(cell, "EEF3F5")
        _set_cell_border(cell)
        for paragraph in cell.paragraphs:
            paragraph.paragraph_format.space_after = Pt(2)
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(10)
    for idx, value in enumerate(values):
        cell = meta.cell(1, idx)
        cell.text = str(value)
        _set_cell_border(cell)
        for paragraph in cell.paragraphs:
            paragraph.paragraph_format.space_after = Pt(0)
            for run in paragraph.runs:
                run.font.size = Pt(10.5)

    summary = doc.add_paragraph()
    summary.paragraph_format.space_before = Pt(18)
    summary.paragraph_format.space_after = Pt(0)
    _add_run(summary, report["headline_de"], size=12, color=INK, bold=True)

    doc.add_page_break()


def _add_bullets(doc: Document, heading: str, rows: list[str]) -> None:
    doc.add_paragraph(heading, style="Heading 1")
    for row in rows:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        _add_run(p, row, size=11, color=INK)


def _add_kpi_table(doc: Document, report: dict[str, Any]) -> None:
    doc.add_paragraph("Turnierbild in Kennzahlen", style="Heading 1")
    table = doc.add_table(rows=1, cols=4)
    table.autofit = False
    widths = [Inches(1.6), Inches(1.6), Inches(1.6), Inches(1.7)]
    for col, width in zip(table.columns, widths):
        col.width = width
    headers = ["Spiele", "Tore", "Tore/Spiel", "Wetterkante"]
    values = [
        str(report["finished_matches"]),
        str(report["total_goals"]),
        str(report["goals_per_match"]),
        f"{report['weather_edge_confirmed']}/{report['comparable_matches']} ({report['weather_edge_hit_rate']}%)",
    ]
    hdr = table.rows[0].cells
    for idx, header in enumerate(headers):
        hdr[idx].text = header
        _set_cell_shading(hdr[idx], "EEF3F5")
        _set_cell_border(hdr[idx])
        for run in hdr[idx].paragraphs[0].runs:
            run.font.bold = True
            run.font.size = Pt(10)
    row = table.add_row().cells
    for idx, value in enumerate(values):
        row[idx].text = value
        _set_cell_border(row[idx])
        for paragraph in row[idx].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.size = Pt(12)
                run.font.bold = True

    doc.add_paragraph(
        "Der Gruppenphasen-Datensatz kombiniert Resultate, Event-Coverage sowie Forecast- und Weather-Fit-Signale. "
        "Ist-Wetter bleibt bewusst ausgeklammert, solange keine belastbaren Messdaten importiert sind."
    )


def _add_featured_matches(doc: Document, report: dict[str, Any]) -> None:
    doc.add_paragraph("Kontextbeispiele", style="Heading 1")
    sections = [
        ("Bestaetigte Wetterkanten", report["featured_matches"]["confirmed"], TEAL),
        ("Verpasste Wetterkanten", report["featured_matches"]["missed"], CORAL),
        ("Remis trotz klarer Kante", report["featured_matches"]["draw"], BLUE),
    ]
    for title, rows, accent in sections:
        doc.add_paragraph(title, style="Heading 2")
        if not rows:
            doc.add_paragraph("In dieser Kategorie liegen im aktuellen Scope keine belastbaren Beispiele vor.")
            continue
        for row in rows:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            _add_run(p, f"{row['match_id']} · {row['label']} · {row['result']}", size=11, color=accent, bold=True)
            body = doc.add_paragraph()
            body.paragraph_format.space_after = Pt(6)
            _add_run(
                body,
                f"Ort: {row.get('host_city') or '–'} | Datum: {row.get('local_date') or '–'} | "
                f"Weather Load: {row.get('weather_load_score') or '–'}/100 | Edge-Gap: {row.get('gap') or '–'}/100.",
                size=10.5,
                color=MUTED,
            )


def _add_team_tables(doc: Document, report: dict[str, Any]) -> None:
    doc.add_paragraph("Teamprofile", style="Heading 1")
    specs = [
        ("Top-Offensiven", report["team_leaders"]["attack"], "goals_for"),
        ("Beste Tordifferenz", report["team_leaders"]["goal_difference"], "goal_difference"),
        ("Meiste Gegentore", report["team_leaders"]["conceded"], "goals_against"),
    ]
    for title, rows, metric_key in specs:
        doc.add_paragraph(title, style="Heading 2")
        table = doc.add_table(rows=1, cols=4)
        table.autofit = False
        widths = [Inches(0.7), Inches(2.6), Inches(1.1), Inches(2.1)]
        for col, width in zip(table.columns, widths):
            col.width = width
        header_cells = table.rows[0].cells
        for idx, text in enumerate(["#", "Team", "Sp", "Wert"]):
            header_cells[idx].text = text
            _set_cell_shading(header_cells[idx], "EEF3F5")
            _set_cell_border(header_cells[idx])
        for index, row in enumerate(rows[:5], start=1):
            cells = table.add_row().cells
            values = [str(index), f"{row['flag']} {row['name_de']}", str(row["played"]), str(row[metric_key])]
            for idx, value in enumerate(values):
                cells[idx].text = value
                _set_cell_border(cells[idx])


def _add_outlook(doc: Document, report: dict[str, Any]) -> None:
    readiness = report["knockout_readiness"]
    coverage = report["event_coverage"]
    doc.add_paragraph("K.o.-Runden-Ausblick", style="Heading 1")
    doc.add_paragraph(
        f"Vor dem weiteren Turnierverlauf sind {readiness['upcoming_matches']} K.o.-Spiele ohne Endstand im Export. "
        f"Davon tragen {readiness['forecast_matches']} bereits Forecasts und {readiness['weather_fit_matches']} belastbare Weather-Fit-Werte."
    )
    doc.add_paragraph(
        f"Die Event-Coverage liegt bei {coverage['goal_event_matches']}/{coverage['finished_matches']} Spielen mit Tor-Events, "
        f"{coverage['lineup_matches']} Spielen mit kompletter Startelf und {coverage['hydration_matches']} Spielen mit Hydration-Markern."
    )
    doc.add_paragraph(
        report["method_note_de"],
    )


def build_docx(data_path: Path = DATA_JS, output_path: Path | None = None) -> Path:
    payload = _load_payload(data_path)
    report = ((payload.get("reports") or {}).get("group_stage_2026")) or {}
    if not report:
        raise ValueError("Missing reports.group_stage_2026 in website/mvp/data.js")

    output = output_path or ROOT / "reports" / "weather-cup-2026-group-stage-report.docx"
    output.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    _style_doc(doc)
    _add_cover(doc, report, payload.get("metadata") or {})
    doc.add_paragraph("Executive Summary", style="Heading 1")
    doc.add_paragraph(report["summary_de"])
    _add_bullets(doc, "Key Findings", report["key_findings_de"])
    _add_kpi_table(doc, report)
    _add_featured_matches(doc, report)
    _add_team_tables(doc, report)
    _add_outlook(doc, report)
    doc.save(output)
    return output


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description="Build group-stage Weather Cup DOCX report")
    parser.add_argument("--data", default=str(DATA_JS))
    parser.add_argument("--output", default=str(ROOT / "reports" / "weather-cup-2026-group-stage-report.docx"))
    args = parser.parse_args(argv)
    path = build_docx(Path(args.data), Path(args.output))
    print(path)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
